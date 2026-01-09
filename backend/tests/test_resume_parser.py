"""
Tests for ResumeParser service
"""
import pytest
from app.services.resume_parser import ResumeParser


@pytest.fixture
def resume_parser():
    """Create ResumeParser instance"""
    return ResumeParser()


@pytest.fixture
def sample_txt_content():
    """Sample TXT resume content"""
    return b"""John Doe
Software Engineer
Email: john.doe@email.com
Phone: (555) 123-4567

EXPERIENCE
Software Engineer | Tech Corp | 2020-Present
- Developed web applications using Python and React
- Led team of 5 developers
- Implemented CI/CD pipelines

EDUCATION
BS Computer Science | State University | 2020

SKILLS
Python, JavaScript, React, Docker, AWS
"""


@pytest.fixture
def sample_pdf_content():
    """Sample PDF content (minimal valid PDF)"""
    # Minimal PDF structure
    return b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >>
endobj
4 0 obj
<< /Length 44 >>
stream
BT /F1 12 Tf 100 700 Td (John Doe Software Engineer) Tj ET
endstream
endobj
xref
0 5
trailer
<< /Size 5 /Root 1 0 R >>
startxref
200
%%EOF"""


class TestResumeParser:
    """Test cases for ResumeParser"""

    @pytest.mark.unit
    async def test_parse_txt_file(self, resume_parser, sample_txt_content):
        """Test parsing TXT file"""
        result = await resume_parser.parse_file(sample_txt_content, "resume.txt")
        
        assert 'raw_text' in result
        assert 'cleaned_text' in result
        assert 'sections' in result
        assert 'skills' in result
        assert 'years_experience' in result
        assert 'contact_info' in result
        assert 'metadata' in result
        
        assert result['metadata']['filename'] == "resume.txt"
        assert result['metadata']['file_type'] == "txt"
        assert len(result['cleaned_text']) > 0

    @pytest.mark.unit
    async def test_parse_unsupported_file(self, resume_parser):
        """Test parsing unsupported file type"""
        with pytest.raises(ValueError, match="Unsupported file type"):
            await resume_parser.parse_file(b"content", "resume.xyz")

    @pytest.mark.unit
    async def test_parse_empty_file(self, resume_parser):
        """Test parsing empty file"""
        result = await resume_parser.parse_file(b"", "empty.txt")
        
        assert result['raw_text'] == ""
        assert len(result['cleaned_text']) == 0
        assert result['metadata']['length'] == 0

    @pytest.mark.unit
    def test_get_resume_summary(self, resume_parser):
        """Test resume summary generation"""
        parsed_resume = {
            'skills': ['Python', 'JavaScript', 'React'],
            'years_experience': 5,
            'sections': {
                'education': 'BS Computer Science'
            }
        }
        
        summary = resume_parser.get_resume_summary(parsed_resume)
        
        assert "5+ years" in summary
        assert "Python" in summary or "Skills" in summary

    @pytest.mark.unit
    def test_get_resume_summary_empty(self, resume_parser):
        """Test summary with minimal data"""
        parsed_resume = {
            'skills': [],
            'years_experience': 0,
            'sections': {}
        }
        
        summary = resume_parser.get_resume_summary(parsed_resume)
        assert "Resume parsed successfully" in summary

    @pytest.mark.integration
    async def test_parse_docx_file(self, resume_parser):
        """Test parsing DOCX file (requires python-docx)"""
        # This test would require creating a sample DOCX file
        # For now, we'll skip if docx module is not available
        try:
            import docx
        except ImportError:
            pytest.skip("python-docx not available")
        
        # Create minimal DOCX content
        from docx import Document
        from io import BytesIO
        
        doc = Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("Software Engineer")
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_content = buffer.getvalue()
        
        result = await resume_parser.parse_file(docx_content, "resume.docx")
        
        assert result['metadata']['file_type'] == "docx"
        assert len(result['raw_text']) > 0
